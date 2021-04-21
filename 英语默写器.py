# -*- coding: utf-8 -*-
"""
Created on Sun Apr 18 20:15:38 2021

英语默写小助手
问题还有：绿色太多（设空句的其它内容变成绿色）、标点更换后大小写出问题、破折号的处理不能这么简单、文件未关闭要处理一下、
         配置文件压根没设置、没了标点没内味
"""
from docx import Document
from docx.shared import RGBColor
import random,re,time,os
  
def change_words(s):#更改标点，但我不会拼
    try:
        s = re.sub("!",".",s)#把叹号换成句号
    except:
        pass
    try:
        s = re.sub("?",".",s)#问号也换掉
    except:
        pass
    try:
        s = re.sub("--"," ",s)#破折号换空格
    except:
        pass
    try:
        s = re.sub("\""," ",s)#引号换空格
    except:
        pass
    try:
        s = re.sub(","," ",s)#逗号换空格
    except:
        pass  
    return s    
def get_s(text_path,FROM=2,TO=5):
    document = Document(text_path)
    s = ''
    for para in document.paragraphs:#读取word文档
        s = s + para.text
    s = change_words(s)#更改标点
    sens = s.split(".")#按句分割
    new_sens = []
    for sen in sens:
        if not sen == ' ' or sen == '':
            new_sens.append(sen)
    answer = new_sens.copy()
    #获取一些幸运句子
    lucky_list = [random.randint(0, len(new_sens)) for _ in range(len(new_sens)//2)]
    SET = set(lucky_list)
    lucky_list = list(SET)#利用集合去重
    lucky_list.sort()#排序
    luck_len_dict = {}
    no_list = []#某些句子长度只有1
    for i in lucky_list:
        sen = new_sens[i]
        luck_len = random.randint(FROM,TO)#挖空长度
        word_list = sen.split(" ")
        LEN = len(word_list)
        S_len = len(sen.split(" "))
        while luck_len > 1 and luck_len > S_len:
            luck_len -= 1#确保挖的空不会超出句子
        if LEN >= 2:
            lucky_word = random.randint(0, LEN-luck_len)#index要少一个
        else:
            no_list.append(i)
            continue
        luck_len_dict[i] = luck_len
        word_list[lucky_word] = "({})".format(luck_len)#括号中间写长度
        for i2 in range(luck_len):
            if not i2 == 0:#不能再把之前的括号删掉
                word_list[lucky_word+i2] = ''#删掉之后的
        sen = ''
        for word in word_list:
            sen = sen + word + " "
        new_sens[i] = sen#处理后的句子再转回去
    for no in no_list:
        lucky_list.append(no)
    s = ''
    for sen in new_sens:
        s = s + sen + "."
    s = re.sub(r"[ ]{2,}"," ",s)#多于一个的空格转为一个空格
    return s,answer,lucky_list,luck_len_dict,lucky_word

def dig_hole(text_path,file_path):#挖空（洞）
    while True:
        try:#从根本上解决出错的问题
            s,answer,lucky_list,luck_len_dict,lucky_word = get_s(text_path)
            break
        except:
            continue#多跑几遍绝对没问题(doge)
    path = file_path
    document = Document()
    document.save(path)
    document = Document(path)#再开一遍，清空内容
    document.add_paragraph(s)
    document.save(path)
    return answer,lucky_list,luck_len_dict,lucky_word

def correct(answer,lucky_list,luck_len_dict,lucky_word,file_path):
    document = Document(file_path)
    s = ''
    for para in document.paragraphs:
        s = s + para.text
    s = change_words(s)#更改标点
    s_list = s.split(".")#同样按句分割答案
    for i in range(len(s_list)):#最后会多一个空格
        list_s = list(s_list[i])
        if len(list_s) >= 1:
            while list_s[-1] == " ":
                list_s.pop()
        _s_ = ''
        for w in list_s:
            _s_ = _s_ + w
        s_list[i] = _s_
    er = 0#error
    al = 0#all
    for key in luck_len_dict:
        al += luck_len_dict[key] + 1#加一才能保证准确
    error_dict = {}#对了用绿色,错了用红色，精确到单词
    correct_list = []#正确的句子用绿色
    for i in lucky_list:
        if s_list[i] == answer[i]:#如果对了就过了
            correct_list.append(i)
            continue
        else:#没对就费事了
            error_dict[i] = []#往字典里加空列表
            error_words = s_list[i].split(" ")
            correct_words = answer[i].split(" ")
            if len(error_words) == len(correct_words):#长度相等
                for j in range(len(error_words)):
                    if not error_words[j] == correct_words[j]:#没对的话：
                        error_dict[i].append(j)#记下应标红的单词，及其所在的句子
            elif len(error_words) != len(correct_words):#不一样的话：
                for j in range(lucky_word,lucky_word+luck_len_dict[i]+1):#可能因lucky_word的位置发生改变而标错词
                    error_dict[i].append(j)
            '''这部分没有用了，哈哈哈
            else:#填多了（为什么会填多了？）
                print(error_words,correct_words)
                f=0
                b=0
                for j in range(len(error_words)):#从前面走一遍
                    if error_words[j] == correct_words[j]:
                        continue
                    else:
                        f = j
                        break
                error_words.reverse()#倒转原答
                correct_words.reverse()#倒转标答
                sdrow_rorre = error_words.copy()
                sdrow_tcerroc = correct_words.copy()
                error_words.reverse()#倒转回原答
                correct_words.reverse()#倒转回标答
                for j in range(len(sdrow_rorre)):#从后面走一遍
                    if sdrow_rorre[j] == sdrow_tcerroc[j]:
                        continue
                    else:
                        b = j
                        break
                b = len(error_words) - b
                print(f,b)
                if f != b:
                    for j in range(f,b+1):
                        print(error_dict[i])
                        error_dict[i].append(j)
                else:
                    error_dict[i].append(f)
            '''
    document = Document()
    document.save(file_path)
    document = Document(file_path)#再开一遍，清空内容
    p = document.add_paragraph()
    for i in range(len(s_list)):
        if i not in lucky_list:
            r = p.add_run(s_list[i]+".")#没有挖空的直接放进去就好了，记得补句点
        elif i in correct_list:
            r = p.add_run(s_list[i]+".")
            r.font.color.rgb = RGBColor(0,255,0)#写对的句子用绿色！
        else:
            for I in range(len(answer[i].split(" "))):
                if I == len(answer[i].split(" "))-1:
                    r = p.add_run(answer[i].split(" ")[I]+".")#补句点
                else:
                    r = p.add_run(answer[i].split(" ")[I]+" ")#补空格
                if not I in error_dict[i]:#往字典里加列表为的就是这个
                    r.font.color.rgb = RGBColor(0,255,0)#正确单词用绿色
                else:
                    r.font.color.rgb = RGBColor(255,0,0)#错误单词用红色
                    er += 1
    document.save(file_path)
    if len(correct_list) == 0 or (al-er)/al <= 0:
        print((al-er)/al)
        return 0
    else:
        return (al-er)/al #返回正确率
def out(accuracy):
    if accuracy == 1:
        print("全对，真棒！")
    elif accuracy == 0:
        print("全错，真有你的")
    else:
        print("改完了，{}分，去看看吧".format(accuracy*100))

def main(text_path="D://英语课文//原文//",file_path="D://英语课文//答卷//"):
    path_list = []
    whole_path_list = []
    for path in os.listdir(text_path):
        T_path = text_path + path
        F_path = file_path + path
        path_list.append(path)
        whole_path_list.append((T_path,F_path))
    while True:
        try:
            for path in path_list:
                print(path)
            index = int(input("请选择文章（序号）：")) - 1#从一开始排序，所以减一
            if index > len(path_list)-1:
                path_list.error()#只是出个错而已
            break
        except:
            print("重写")
    T_path,F_path = whole_path_list[index]
    while True:
        try:
            answer,lucky_list,luck_len_dict,lucky_word = dig_hole(text_path=T_path,file_path=F_path)
            break
        except:
            print("\r文件没有关",end='')
            time.sleep(0.5)
    while True:
        try:
            accuracy = input("开始做吧,做好了就随便打些东西进来：")#用后面出现的变量是为了不报错且不影响
            accuracy = correct(answer,lucky_list,luck_len_dict,lucky_word,file_path=F_path)
            break
        except PermissionError:
            print("文件没有关")
    out(accuracy)
    
yes = ''
while yes == '':
    main()
    yes = input("还继续吗？继续就回车，否则按任意键后回车")
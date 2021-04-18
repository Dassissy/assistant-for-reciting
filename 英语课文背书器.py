# -*- coding: utf-8 -*-
"""
Created on Sun Apr 18 11:57:15 2021

英语背书挖空小助手
"""
from docx import Document
import random,re

def get_s():
    document = Document("D://英语课文//1.docx")
    s = ''
    for para in document.paragraphs:#读取word文档
        s = s + para.text
    try:
        s = re.sub("!",".",s)#把叹号换成句号
    except:
        pass
    try:
        s = re.sub("?",".",s)#问号也换掉
    except:
        pass
    sens = s.split(".")#按句分割
    new_sens = []
    for sen in sens:
        if not sen == '':
            new_sens.append(sen)
    lucky_list = [random.randint(0, len(new_sens)) for _ in range(len(new_sens)//3)]#取一些幸运句子
    SET = set(lucky_list)
    lucky_list = list(SET)#利用集合去重
    lucky_list.sort()#排序
    for i in lucky_list:
        sen = new_sens[i]
        try:
            sen = re.sub(","," ",sen)
        except:
            pass
        luck_len = random.randint(2, 5)#挖空长度2-5
        word_list = sen.split(" ")
        LEN = len(word_list)
        while luck_len > 1:
            if luck_len > len(sen.split(" ")):
                luck_len -= 1#确保挖的空不会超出句子
            else:
                break
        try:
            lucky_word = random.randint(0, LEN-luck_len-1)#index要少一个
        except:
            continue
        word_list[lucky_word] = "({})".format(luck_len)#括号中间写长度
        for i2 in range(luck_len):
            if not i2 == 0:#不能再把之前的括号删掉
                word_list[lucky_word+i2] = ''#删掉之后的
        sen = ''
        for word in word_list:
            sen = sen + word + " "
        new_sens[i] = sen#处理后的句子再转回去
    s = ''
    for sen in new_sens:
        s = s + sen + "."
    return s

def main():
    while True:
        try:#从根本上解决出错的问题
            s = get_s()
            break
        except:
            continue#多跑几遍绝对没问题(doge)
    path = "D://英语课文//2.docx"
    try:
        document = Document(path)#要是存在此文档的话就直接开
        document.clear()#清空，不然越跑越多
    except:
        document = Document()
        document.save(path)#不存在就建一个
        document = Document(path)#然后再开
    document.add_paragraph(s)
    document.save(path)

main()
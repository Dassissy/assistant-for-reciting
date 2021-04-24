# -*- coding: utf-8 -*-
"""
Created on Fri Apr 23 22:45:21 2021

语文默写小助手
"""
from docx import Document
from docx.shared import RGBColor
import random,re,time,os
  
def change_words(s):#更改标点，但我不会拼
    try:
        s = re.sub("[!?！？.]","。",s)#叹号、问号、省略号换句号
    except:
        pass
    try:
        s = re.sub("[——-\"“”]","",s)#破折号、引号弄没掉
    except:
        pass
    try:
        s = re.sub("[;；]","，",s)#分号换逗号
    except:
        pass
    try:
        s = re.sub("[ (\u3000)]","",s)#空格消失术
    except:
        pass
    #顿号没有刻意处理，试试行不行
    return s   

def get_s(text_path,FROM=2,TO=5,num_of_Qs=2):
    document = Document(text_path)
    s = ''
    for para in document.paragraphs:#读取word文档
        s = s + para.text + "\n"
    s = change_words(s)#更改标点
    sens = s.split("。")#按句分割
    new_sens = []
    for sen in sens:#此处不能用if not sen == " " or not sen == "": new_sens.append(sen)的结构
        if sen == ' ' or sen == '':
            pass
        else:
            new_sens.append(sen)
    #获取一些幸运句子
    #len返回的值是从1开始数的，随机数是含末尾的（len(new_sens)-1)
    #print(("new_sens and its len are:{},{}").format(new_sens,len(new_sens)))
    lucky_list = [random.randint(0, len(new_sens)-1) for _ in range(len(new_sens)//num_of_Qs)]
    SET = set(lucky_list)
    lucky_list = list(SET)#利用集合去重
    lucky_list.sort()#排序
    luck_len_dict = {}
    no_list = []#某些句子长度只有1
    answer = new_sens.copy()
    #print("answer is:{}".format(answer))
    lucky_word_list = []#记录lucky_word
    for i in lucky_list:
        #print(("i and lucky_list are:{},{}").format(i,lucky_list))
        sen = new_sens[i]
        luck_len = random.randint(FROM,TO)#挖空长度
        word_list = sen.split("，")#以以逗号隔开的小句为单位考察
        #print(("word_list is:{}").format(word_list))
        LEN = len(word_list)
        #print(("LEN is:{}").format(LEN))
        while luck_len > 1 and luck_len > LEN:
            #print(("i and luck_len are:{},{}").format(i,luck_len))
            luck_len -= 1#确保挖的空不会超出句子
        if LEN >= FROM:#若到最后LEN仍然在FROM以上
            lucky_word = random.randint(0, LEN-luck_len)#index要少一个，但len()会多一个，所以最后没有+1也没有-1
            lucky_word_list.append(lucky_word)
        else:
            no_list.append(i)
            continue
        luck_len_dict[i] = luck_len
        #print(("luck_len is:{}").format(luck_len))
        word_list[lucky_word] = "({})".format(luck_len)#括号中间写长度
        for i2 in range(luck_len):
            if not i2 == 0:#不能再把之前的括号删掉
                word_list[lucky_word+i2] = ''#删掉之后的
        sen = ''
        for word in word_list:
            if not word_list.index(word) == len(word_list)-1:
                sen = sen + word + "，"
            else:
                sen = sen + word
        new_sens[i] = sen#处理后的句子再转回去
    for no in no_list:
        lucky_list.remove(no)
    s = ''
    for sen in new_sens[:-1]:#这最后一个不知为何是空字符串
        s = s + sen + "。"
    #print(("luck_len_dict,lucky_word_list are:{},{}").format(luck_len_dict,lucky_word_list))
    s = re.sub(r"[ ]*","",s)#空格消失术
    s = modify_s(s)#对s的内容略做修改
    return s,answer,lucky_list,luck_len_dict,lucky_word_list

def modify_s(s):
    s = re.sub("。","！。",s)#___。___，。___。 -> ___！。___，！。___！。
    s_list = s.split("。")[:-1]# -> ___！___，！___！#最后会有一个空字符串，不要它
    #print(("s_list before modifying is:{}").format(s_list))
    for sen in s_list:
        if re.search(r"[(][1234567890]*?[)]", sen):#若出现"(数字)"的结构
            if re.search(r"，，！",sen):#若出现"，，！"，那么：#可能只有两个空
                new_sen = re.sub(r"，！","！",sen)# -> ___！___！___！
            else:
                new_sen = sen
            new_sen = re.sub(r"[(][1234567890]*?[)]","",new_sen)#不论如何，都将"(数字)"结构删除
            s_list[s_list.index(sen)] = new_sen[:-1]#最后的感叹号就不要了
        else:
            continue
    #print(("s_list after modifying is:{}").format(s_list))
    s = ''
    for S in s_list:
        s = s + S + "。"
    s = re.sub("！","",s)
    #print(("s is:{}").format(s))
    return s

def dig_hole(text_path,file_path,FROM,TO,num_of_Qs):#挖空（洞）
    s,answer,lucky_list,luck_len_dict,lucky_word_list = get_s(text_path,FROM,TO,num_of_Qs)
    lucky_list = set(lucky_list)
    lucky_list = list(lucky_list)#不知为何这里仍有重复，于是再去重一次
    #print("lucky_list after duplicate removal is:{}".format(lucky_list))
    path = file_path
    document = Document()
    document.save(path)
    document = Document(path)#再开一遍，清空内容
    document.add_paragraph(s)
    document.save(path)
    return answer,lucky_list,luck_len_dict,lucky_word_list,s

def correct(answer,lucky_list,luck_len_dict,lucky_word_list,s,file_path):
    document = Document(file_path)
    s = ''
    for para in document.paragraphs:
        s = s + para.text
    s = change_words(s)#更改标点
    s_list = s.split("。")#同样按句分割答案
    for i in range(len(s_list)):#最后会多一个空格
        list_s = list(s_list[i])
        if len(list_s) >= 1:
            while list_s[-1] == " ":
                list_s.pop()
        _s_ = ''
        for w in list_s:
            _s_ = _s_ + w
        s_list[i] = _s_
    er = 0#error
    al = 0#all
    for key in luck_len_dict:
        al += luck_len_dict[key]#不要+1！！！
    error_dict = {}#错了用红色，精确到字{i1:{j1:[k1,k2],j2:[k3,k4,k5]},i2:{j3:[k6,k7],j4:[k8]}}
    correct_list = []#正确的句子用绿色
    ilist = []
    for i in range(len(lucky_list)):
        if lucky_list[i] not in luck_len_dict:#弹出放弃的句子（过短）
            ilist.append(i)
    #print(("ilist is:{}\nlucky_list is:{}\nluck_len_dict is:{}").format(ilist,lucky_list,luck_len_dict))
    count = 0#使用count抵消掉弹出后列表长度的缩减，这里又可以用了
    for i in ilist:
        lucky_list.pop(i-count)
        count += 1
    for i in range(len(lucky_list)):
        lucky_list[i] = (lucky_list[i],i)
    #print(("lucky_list with count is:{}").format(lucky_list))
    for i,count in lucky_list:
        if s_list[i] == answer[i]:#如果对了就过了
            correct_list.append(i)
            continue
        else:#没对就费事了
            error_dict[i] = {}#往字典里加空字典
            error_words = s_list[i].split("，")
            correct_words = answer[i].split("，")
            if len(error_words) == len(correct_words):#长度相等
                for j in range(len(correct_words)):
                    if not error_words[j] == correct_words[j]:#没对的话：)
                        #print(("error_words[j] and correct_words[i] are:{},{}").format(error_words[j],correct_words[j]))
                        error_zi = list(error_words[j])
                        correct_zi = list(correct_words[j])
                        #print(("error_zi and correct_zi are:{},{}").format(error_zi,correct_zi))
                        if len(error_zi) == len(correct_zi):#梅开二度按字数来分
                            count = 0
                            for k in range(len(correct_zi)):
                                if not correct_zi[k] == error_zi[k]:
                                    if count == 0:
                                        error_dict[i][j] = [k]#创建列表
                                        count += 1
                                    else:
                                        error_dict[i][j].append(k)#往列表内追加元素
                        else:#字数不同则全部放进去
                            error_dict[i][j] = []
                            for k in range(len(correct_zi)):
                                error_dict[i][j].append(k)
            else:#不一样的话：
                """A = 1
                try:
                    print(("lucky_word_list[count],luck_len_dict[i] are:{},{}").foramt(lucky_word_list[count],luck_len_dict[i]))
                except:
                    print(("lucky_word_list,luck_len_dict,i,count are:{},{},{},{}").format(lucky_word_list,luck_len_dict,i,count))
                    pass"""
                for j in range(lucky_word_list[count],lucky_word_list[count]+luck_len_dict[i]):
                    #a = (lucky_word_list[count],lucky_word_list[count]+luck_len_dict[i])
                    error_dict[i][j] = []
                    print(("range and correct_zi are:{},{}").format((lucky_word_list[count],lucky_word_list[count]+luck_len_dict[i]-1),list(correct_words[j])))
                    for k in range(len(list(correct_words[j]))):
                        error_dict[i][j].append(k)
    #print(("error_dict is:{}").format(error_dict))
    document = Document()
    document.save(file_path)
    document = Document(file_path)#再开一遍，清空内容
    p = document.add_paragraph()
    while s_list[-1] == "":
        s_list.pop()
    for i in range(len(lucky_list)):
        lucky_list[i] = lucky_list[i][0]#又要从元组变回单个的数字
    for i in range(len(answer)):
        #print(("i and lukcy_list are:{} and {}:").format(i,lucky_list))
        if i not in lucky_list:
            """
            print(("i:{} is not in lucky_list").format(i))
            try:
                print(("answer[i] is:{}").format(answer[i]))
            except:
                print(("s_list[i] is:{}").format(s_list[i]))
                """
            r = p.add_run(answer[i]+"。")#没有挖空的直接放进去就好了，记得补句号
        elif i in correct_list:
            #b = "sign"
            #print(("i,correct_list are:{},{}").format(i,correct_list))
            r = p.add_run(answer[i]+"。")#最后会多一个句号，所以这里可以改一改
            r.font.color.rgb = RGBColor(0,255,0)#写对的句子用绿色！
        else:
            for I in range(len(answer[i].split("，"))):
                #print(("i and error_dict are:{} and\n{}").format(i,error_dict))
                #print(("error_dict[i] is:{}").format(error_dict[i]))
                if not I in error_dict[i]:#往字典里加字典为的就是这个
                    if I == len(answer[i])-1:
                        r = p.add_run(answer[i].split("，")[I]+"。")
                    else:
                        r = p.add_run(answer[i].split("，")[I]+"，")
                    r.font.color.rgb = RGBColor(0,255,0)#正确的小段用绿色
                else:
                    er += 1#算分按小段算,故er先加上1
                    #print(("answer[i] and I are:{},{}").format(answer[i],I))
                    ans_zi = list(answer[i].split("，")[I])
                    for II in range(len(ans_zi)):
                        if II == len(ans_zi)-1:#最后一个字后是句号或逗号
                            if I == len(answer[i].split("，"))-1:
                                #print(("now I:{} == len(answer[i].split('，'))-1").format(I))
                                #print(("answer[i].split('，') is:{}").format(answer[i].split('，')))
                                r = p.add_run(ans_zi[II]+"。")
                            else:
                                #print(("now I:{} != len(answer[i].split('，'))-1").format(I))
                                #print(("answer[i].split('，') is:{}").format(answer[i].split('，')))
                                r = p.add_run(ans_zi[II]+"，")
                        else:
                            r = p.add_run(ans_zi[II])#小段中间的字就什么都不需要了
                        if not II in error_dict[i][I]:
                            r.font.color.rgb = RGBColor(0,255,0)#正确的字用绿色
                        else:
                            r.font.color.rgb = RGBColor(255,0,0)#错字用红色
    document.save(file_path)
    #print(("er,al are:{},{}".format(er,al)))
    accuracy = (al-er)/al
    if accuracy == 0:
        return 0,0
    else:
        #wao = str(answer)+"\n"+str(lucky_list)+"\n"+str(error_dict)+"\n"+str(error_words)
        #print("they are:answer, lucky_list,error_dict,error_words\n{}".format(wao))
        return accuracy,al-er #返回正确率及正确空数
def out(accuracy,correct_num):
    if accuracy == 1:
        print("全对，真棒！")
    elif accuracy == 0:
        print("全错，真有你的")
    else:
        print("改完了，正确{}个空，计{}分，去看看吧".format(correct_num,round(accuracy*100, 1)))#没有对精度的要求，故直接使用round()

def make_path(path):
    path_list = path.split("//")
    path_list.pop()#最后还有一个元素是空字符串，不要在上面直接pop()，此函数返回被弹出的元素
    #print(("path_list is:{}").format(path_list))
    path_now = ''
    c = 0
    #print(("new path_now is:{}").format(path_now))
    while c+1 <= len(path_list):
        path_now = path_now + path_list[c] + "//"
        #print(("path_now is:{}").format(path_now))
        if not os.path.exists(path_now):
            #print(("not exist path:{}").format(path_now))
            os.mkdir(path_now)
        c += 1
    #print(("path_now is:{}").format(path_now))

def read_config():#读取配置
    try:
        config_path = os.getcwd() + "\\" + "config\\Chinese_config.txt"
        #print(("config_path is:{}").format(config_path))
        with open(config_path,"r",encoding="utf-8") as f:#把相对位置转换为绝对位置
            lines = f.readlines()
            new_lines = []
            for line in lines:
                line = line.split("：")[1]
                new_lines.append(line)
            #最后有一个换行符所以全部取到倒数第二个字符
            FROM = new_lines[0][:-1]
            TO = new_lines[1][:-1]
            text_path = new_lines[2][:-1]
            file_path = new_lines[3][:-1]
            num_of_Qs = new_lines[4][:-1]
            #print(("FROM,TO,num_of_Qs are:{},{},{}").format(FROM,TO,num_of_Qs))
            FROM,TO,num_of_Qs = int(FROM),int(TO),int(num_of_Qs)#这几个是int类型的
            #print(("file_path is:{}").format(file_path))
            make_path(text_path)#建立问卷路径，集中错误：没有弄原文
            make_path(file_path)#建立答卷路径
    except:
        print("配置文件出错")
        return "again now"
    return FROM,TO,text_path,file_path,num_of_Qs
def main(text_path="D://英语课文//原文//",file_path="D://英语课文//答卷//"):
    try:
        FROM,TO,text_path,file_path,num_of_Qs = read_config()#读取配置
    except ValueError:
        return "again now"
    path_list = []
    whole_path_list = []
    for path in os.listdir(text_path):
        T_path = text_path + path
        F_path = file_path + path
        path_list.append(path)
        whole_path_list.append((T_path,F_path))
    while True:
        if not path_list:
            print("你没有设置原文")
            return "again now"
        try:
            for path in path_list:
                print(path)
            index = int(input("请选择文章（序号）：")) - 1#从一开始排序，所以减一
            if index > len(path_list)-1:#len()返回的值比索引大1（重要的东西写三遍）
                path_list.error()#只是出个错而已
            break
        except:
            print("重写")
    T_path,F_path = whole_path_list[index]
    try:
        Document(T_path)
    except:
        print("这是一个空文档或不是docx文档")
        return "again now"
    while True:
        try:
            answer,lucky_list,luck_len_dict,lucky_word,s = dig_hole(text_path=T_path,file_path=F_path,FROM=FROM,TO=TO,num_of_Qs=num_of_Qs)
            #print(("luck_len_dict is:{}").format(luck_len_dict))
            if not luck_len_dict:#如果没有任何有长的空，那么这传进来的东西就是没有实质内容了
                print("此文章无有效句子")
                return "again now"
            break
        except PermissionError:
            print("\r文件没有关",end='')
        time.sleep(0.5)
    al = 0#all
    for key in luck_len_dict:
        al += luck_len_dict[key]#不要+1！！！
    while True:
        try:
            
            accuracy = input(("本次小测共设空{}处，按任意键开始批改：").format(al))#用后面出现的变量是为了不报错且不影响
            accuracy,correct_num = correct(answer,lucky_list,luck_len_dict,lucky_word,s,file_path=F_path)
            break
        except PermissionError:
            print("文件没有关")
    out(accuracy,correct_num)
    
yes = ''
while yes == '':
    main()
    yes = input("还继续吗？继续就回车，否则按任意键后回车")
